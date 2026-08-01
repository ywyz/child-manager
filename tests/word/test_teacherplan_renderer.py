"""T130 固定 teacherplan.docx 渲染结构与样式 RED。"""

import json
from copy import deepcopy
from hashlib import sha256
from importlib import import_module
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

TEMPLATE = Path("templates/teacherplan/teacherplan.docx")
FIXTURE = Path("tests/fixtures/word/daily_activity_plan_v1.json")
TEMPLATE_SHA256 = "72ee26e7cb8f510a11bc303b7a967c2a375fe436b5c8a72822ee9ccbfe235043"


def _renderer_type() -> type[Any]:
    try:
        module = import_module("packages.backend.integrations.files.teacherplan_renderer")
    except ModuleNotFoundError:
        pytest.fail("T130 TeacherplanRenderer 尚未实现")
    renderer = getattr(module, "TeacherplanRenderer", None)
    if renderer is None:
        pytest.fail("T130 TeacherplanRenderer 尚未实现")
    return renderer


def _fixture() -> dict[str, Any]:
    return json.loads(FIXTURE.read_text(encoding="utf-8"))


def _render(sample: dict[str, Any] | None = None) -> tuple[bytes, Any]:
    payload = sample or _fixture()
    renderer = _renderer_type()(TEMPLATE, expected_sha256=TEMPLATE_SHA256)
    rendered = renderer.render(
        context_snapshot=payload["context_snapshot"],
        content_snapshot=payload["content_snapshot"],
    )
    assert isinstance(rendered, bytes)
    return rendered, Document(BytesIO(rendered))


def _east_asia_font(run: Any) -> str | None:
    if run._element.rPr is None or run._element.rPr.rFonts is None:
        return None
    return run._element.rPr.rFonts.get(qn("w:eastAsia"))


def test_renderer_uses_copy_and_preserves_template_structure_fonts_and_fixed_labels() -> None:
    before = sha256(TEMPLATE.read_bytes()).hexdigest()
    rendered, document = _render()

    assert rendered
    assert sha256(TEMPLATE.read_bytes()).hexdigest() == before == TEMPLATE_SHA256
    assert len(document.tables) == 1
    table = document.tables[0]
    assert len(table.rows) == 19 and len(table.columns) == 2
    assert document.paragraphs[0].text == "星河幼儿园一日活动计划（2026.2-2026.6）"
    assert document.paragraphs[1].text == "向日葵班 陈老师 李老师"
    assert table.cell(0, 0).text == "第（四）周"
    assert table.cell(1, 0).text == "周（一）3月2日"
    assert table.cell(2, 0).text == "晨间活动："
    assert "体能大循环" in table.cell(2, 1).text
    assert table.cell(4, 0).text == "晨间谈话："
    assert table.cell(6, 1).text == "活动主题：《寻找春天》"
    assert table.cell(12, 1).text == "游戏区域：建构区、阅读区"
    assert table.cell(15, 1).text == "游戏区域：球类区、平衡区"

    title_run = document.paragraphs[0].runs[0]
    body_run = table.cell(2, 1).paragraphs[0].runs[0]
    assert title_run.font.size is not None and title_run.font.size.pt == 16
    assert title_run.font.name == _east_asia_font(title_run) == "楷体"
    assert body_run.font.size is not None and body_run.font.size.pt == 12
    assert body_run.font.name == _east_asia_font(body_run) == "仿宋"


def test_renderer_preserves_every_template_cell_paragraph() -> None:
    template = Document(str(TEMPLATE))
    _rendered, rendered = _render()
    template_table = template.tables[0]
    rendered_table = rendered.tables[0]

    for row_index in range(len(template_table.rows)):
        for column_index in range(len(template_table.columns)):
            before = template_table.cell(row_index, column_index).paragraphs
            after = rendered_table.cell(row_index, column_index).paragraphs
            assert len(after) == len(before), (row_index, column_index)
            assert [
                paragraph._p.pPr.xml if paragraph._p.pPr is not None else None
                for paragraph in after
            ] == [
                paragraph._p.pPr.xml if paragraph._p.pPr is not None else None
                for paragraph in before
            ]


def test_renderer_preserves_template_header_paragraph_properties() -> None:
    template = Document(str(TEMPLATE))
    _rendered, rendered = _render()

    for paragraph_index in (0, 1):
        before = template.paragraphs[paragraph_index]
        after = rendered.paragraphs[paragraph_index]
        assert after.alignment == before.alignment
        assert before._p.pPr is not None
        assert after._p.pPr is not None
        assert after._p.pPr.xml == before._p.pPr.xml


def test_renderer_preserves_template_emphasis_for_fixed_area_labels() -> None:
    _rendered, document = _render()
    table = document.tables[0]

    for row_index, paragraph_indexes in {
        13: (1, 5),
        14: (0,),
        16: (1, 5),
        17: (0,),
    }.items():
        for paragraph_index in paragraph_indexes:
            assert table.cell(row_index, 1).paragraphs[paragraph_index].runs[0].bold is True


def test_only_ai_added_group_activity_step_is_red() -> None:
    _rendered, document = _render()
    process_cell = document.tables[0].cell(11, 1)
    red_text = "".join(
        run.text
        for paragraph in process_cell.paragraphs
        for run in paragraph.runs
        if run.font.color.rgb == RGBColor(255, 0, 0)
    )
    non_red_text = "".join(
        run.text
        for paragraph in process_cell.paragraphs
        for run in paragraph.runs
        if run.font.color.rgb != RGBColor(255, 0, 0)
    )

    assert "二、新增延伸" in red_text
    assert "请把发现画下来" in red_text
    assert "一、观察图片" in non_red_text
    assert "AI 新增" not in process_cell.text


def test_empty_week_and_reflection_keep_fixed_positions_and_three_rows() -> None:
    sample = deepcopy(_fixture())
    sample["context_snapshot"]["teaching_week_number"] = None
    sample["context_snapshot"]["teaching_week_text"] = None
    _rendered, document = _render(sample)

    assert document.tables[0].cell(0, 0).text == ""
    reflection = document.tables[0].cell(18, 1).text.splitlines()
    assert reflection == ["活动亮点：", "存在问题：", "调整策略："]


def test_missing_wrong_hash_or_wrong_structure_fails_without_rebuilding(tmp_path: Path) -> None:
    renderer_type = _renderer_type()
    with pytest.raises(Exception, match="模板"):
        renderer_type(tmp_path / "missing.docx", expected_sha256=TEMPLATE_SHA256)
    with pytest.raises(Exception, match="哈希"):
        renderer_type(TEMPLATE, expected_sha256="0" * 64)

    malformed = tmp_path / "malformed.docx"
    Document().save(str(malformed))
    actual_hash = sha256(malformed.read_bytes()).hexdigest()
    renderer = renderer_type(malformed, expected_sha256=actual_hash)
    with pytest.raises(Exception, match="结构"):
        renderer.render(
            context_snapshot=_fixture()["context_snapshot"],
            content_snapshot=_fixture()["content_snapshot"],
        )
