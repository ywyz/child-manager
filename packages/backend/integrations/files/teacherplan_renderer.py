"""固定 teacherplan.docx 副本渲染器。"""

from __future__ import annotations

from collections.abc import Iterable, Mapping, Sequence
from copy import deepcopy
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell
from docx.text.paragraph import Paragraph


class TeacherplanTemplateError(ValueError):
    """固定模板缺失、哈希漂移或结构不匹配。"""


class TeacherplanRenderer:
    """只读取固定模板，并在内存副本中替换已确认字段。"""

    def __init__(self, template_path: Path, *, expected_sha256: str) -> None:
        self._template_path = Path(template_path)
        if not self._template_path.is_file():
            raise TeacherplanTemplateError("Word 模板不存在")
        self._template_bytes = self._template_path.read_bytes()
        actual_sha256 = sha256(self._template_bytes).hexdigest()
        if actual_sha256 != expected_sha256:
            raise TeacherplanTemplateError("Word 模板哈希不匹配")

    def render(
        self,
        *,
        context_snapshot: Mapping[str, Any],
        content_snapshot: Mapping[str, Any],
    ) -> bytes:
        document = Document(BytesIO(self._template_bytes))
        self._validate_structure(document)
        self._render_header(document, context_snapshot)
        self._render_table(document, context_snapshot, content_snapshot)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _validate_structure(document: DocumentType) -> None:
        if len(document.paragraphs) < 2 or len(document.tables) != 1:
            raise TeacherplanTemplateError("Word 模板结构不匹配")
        table = document.tables[0]
        if len(table.rows) != 19 or len(table.columns) != 2:
            raise TeacherplanTemplateError("Word 模板结构不匹配")
        expected_labels = {
            2: "晨间活动：",
            4: "晨间谈话：",
        }
        if any(table.cell(row, 0).text != label for row, label in expected_labels.items()):
            raise TeacherplanTemplateError("Word 模板固定文本结构不匹配")

    def _render_header(
        self,
        document: DocumentType,
        context: Mapping[str, Any],
    ) -> None:
        semester_range = self._semester_month_range(
            context.get("semester_start_date"),
            context.get("semester_end_date"),
        )
        self._replace_paragraph(
            document.paragraphs[0],
            [(f"{context.get('kindergarten_name', '')}一日活动计划（{semester_range}）", False)],
            font_name="楷体",
            font_size=16,
        )
        authors = context.get("authors", [])
        author_names = [
            str(author.get("display_name_snapshot", ""))
            for author in authors
            if isinstance(author, Mapping) and author.get("display_name_snapshot")
        ]
        subtitle = " ".join([str(context.get("class_name", "")), *author_names]).strip()
        self._replace_paragraph(
            document.paragraphs[1],
            [(subtitle, False)],
            font_name="楷体",
            font_size=16,
        )

    def _render_table(
        self,
        document: DocumentType,
        context: Mapping[str, Any],
        content: Mapping[str, Any],
    ) -> None:
        table = document.tables[0]
        self._set_cell(table.cell(0, 0), str(context.get("teaching_week_text") or ""))
        self._set_cell(table.cell(1, 0), str(context.get("activity_date_text") or ""))

        morning = self._mapping(content.get("morning_activity"))
        self._set_cell(
            table.cell(2, 1),
            "\n".join(
                (
                    str(morning.get("physical_cycle") or "体能大循环"),
                    f"集体游戏：《{morning.get('group_game', '')}》",
                    f"自主游戏：《{morning.get('free_game', '')}》",
                )
            ),
        )
        self._set_cell(
            table.cell(3, 1),
            self._guidance_text(morning),
        )

        talk = self._mapping(content.get("morning_talk"))
        self._set_cell(table.cell(4, 1), f"话题：《{talk.get('topic', '')}》")
        self._set_cell(
            table.cell(5, 1),
            "问题设计：\n" + self._numbered(talk.get("questions")),
        )

        group = self._mapping(content.get("group_activity"))
        self._set_cell(table.cell(6, 1), f"活动主题：《{group.get('theme', '')}》")
        self._set_cell(table.cell(7, 1), "活动目标：\n" + self._numbered(group.get("objectives")))
        self._set_cell(table.cell(8, 1), "活动准备：\n" + self._numbered(group.get("preparation")))
        self._set_cell(table.cell(9, 1), f"活动重点：{group.get('focus', '')}")
        self._set_cell(table.cell(10, 1), f"活动难点：{group.get('difficulty', '')}")
        self._set_group_process(table.cell(11, 1), group.get("process"))

        self._render_area(table, 12, self._mapping(content.get("indoor_area_game")))
        self._render_area(table, 15, self._mapping(content.get("afternoon_outdoor_game")))

        reflection = self._mapping(content.get("daily_reflection"))
        self._set_cell(
            table.cell(18, 1),
            "\n".join(
                (
                    f"活动亮点：{reflection.get('highlights', '')}",
                    f"存在问题：{reflection.get('issues', '')}",
                    f"调整策略：{reflection.get('adjustments', '')}",
                )
            ),
        )

    def _render_area(self, table: Any, start_row: int, area: Mapping[str, Any]) -> None:
        areas = area.get("areas")
        names = "、".join(str(value) for value in areas) if isinstance(areas, list) else ""
        self._set_cell(table.cell(start_row, 1), f"游戏区域：{names}")
        self._set_cell(table.cell(start_row + 1, 1), self._guidance_text(area))
        self._set_cell(
            table.cell(start_row + 2, 1),
            "支持策略：\n" + self._numbered(area.get("support_strategies")),
        )

    def _set_group_process(self, cell: _Cell, value: object) -> None:
        segments: list[tuple[str, bool]] = [("活动过程：", False)]
        if isinstance(value, list):
            for step in value:
                if not isinstance(step, Mapping):
                    continue
                text = "\n" + str(step.get("heading", ""))
                lines = step.get("lines")
                if isinstance(lines, list):
                    text += "".join(f"\n{index}.{line}" for index, line in enumerate(lines, 1))
                segments.append((text, step.get("is_ai_added") is True))
        self._set_cell(cell, segments=segments)

    @staticmethod
    def _guidance_text(value: Mapping[str, Any]) -> str:
        parts = [
            f"重点指导：{value.get('focus_guidance', '')}",
            "活动目标：",
            TeacherplanRenderer._numbered(value.get("objectives")),
            "指导要点：",
            TeacherplanRenderer._numbered(value.get("guidance_points")),
        ]
        return "\n".join(parts)

    @staticmethod
    def _numbered(value: object) -> str:
        if not isinstance(value, list):
            return ""
        return "\n".join(f"{index}.{item}" for index, item in enumerate(value, 1))

    @staticmethod
    def _mapping(value: object) -> Mapping[str, Any]:
        return value if isinstance(value, Mapping) else {}

    @staticmethod
    def _semester_month_range(start: object, end: object) -> str:
        def month(value: object) -> str:
            parts = str(value).split("-")
            if len(parts) < 2:
                return str(value)
            return f"{int(parts[0])}.{int(parts[1])}"

        return f"{month(start)}-{month(end)}"

    def _set_cell(
        self,
        cell: _Cell,
        text: str | None = None,
        *,
        segments: Sequence[tuple[str, bool]] | None = None,
    ) -> None:
        paragraphs = cell.paragraphs
        lines = self._segment_lines(segments or [(text or "", False)])
        for index, paragraph in enumerate(paragraphs):
            if index == len(paragraphs) - 1:
                assigned = self._join_lines(lines[index:])
            else:
                assigned = lines[index] if index < len(lines) else []
            self._replace_paragraph(
                paragraph,
                assigned or [("", False)],
                font_name="仿宋",
                font_size=12,
            )

    @staticmethod
    def _segment_lines(
        segments: Sequence[tuple[str, bool]],
    ) -> list[list[tuple[str, bool]]]:
        lines: list[list[tuple[str, bool]]] = [[]]
        for text, is_red in segments:
            for index, piece in enumerate(text.split("\n")):
                if index:
                    lines.append([])
                if piece:
                    lines[-1].append((piece, is_red))
        return lines

    @staticmethod
    def _join_lines(
        lines: Sequence[Sequence[tuple[str, bool]]],
    ) -> list[tuple[str, bool]]:
        joined: list[tuple[str, bool]] = []
        for index, line in enumerate(lines):
            if index:
                joined.append(("\n", False))
            joined.extend(line)
        return joined

    @staticmethod
    def _replace_paragraph(
        paragraph: Paragraph,
        segments: Iterable[tuple[str, bool]],
        *,
        font_name: str,
        font_size: int,
    ) -> None:
        run_properties = (
            deepcopy(paragraph.runs[0]._r.rPr)
            if paragraph.runs and paragraph.runs[0]._r.rPr is not None
            else None
        )
        for run in list(paragraph.runs):
            paragraph._p.remove(run._r)
        for text, is_red in segments:
            run = paragraph.add_run(text)
            if run_properties is not None:
                run._r.insert(0, deepcopy(run_properties))
            run.font.name = font_name
            run_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
            run_fonts.set(qn("w:eastAsia"), font_name)
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(255, 0, 0) if is_red else RGBColor(0, 0, 0)
